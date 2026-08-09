#!/usr/bin/env python3
"""
File Renamer & Organizer

A macOS Python application to rename and organize local files using LLM-powered analysis.
Supports multiple LLM providers via LangChain.

Usage:
    uv run renamer.py [OPTIONS] [DIRECTORY]
"""

from __future__ import annotations

import base64
import json
import logging
import os
import re
import subprocess
import tempfile
import time
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Annotated, Iterator, Literal

import typer
import yaml
from langchain_core.language_models import BaseChatModel
from langchain_core.messages import HumanMessage
from pydantic import BaseModel
from rich.console import Console
from rich.panel import Panel
from rich.progress import BarColumn, Progress, SpinnerColumn, TaskProgressColumn, TextColumn
from rich.table import Table
from send2trash import send2trash

# Initialize Rich console and Typer app
console = Console()
app = typer.Typer(
    help="File renamer and organizer using local LLM analysis.",
    no_args_is_help=False,
)

# Configure logging - suppress noisy third-party loggers
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)

# Suppress HTTP request logs from httpx (used by Gemini client)
logging.getLogger("httpx").setLevel(logging.WARNING)
# Suppress Google GenAI AFC (Automatic Function Calling) logs
logging.getLogger("google.genai").setLevel(logging.WARNING)
logging.getLogger("google.generativeai").setLevel(logging.WARNING)
logging.getLogger("google_genai.models").setLevel(logging.WARNING)

# PEP 723 inline-metadata markers for the apply/revert scripts we generate.
# Assembled at runtime rather than written literally: uv scans a script's raw
# text for these markers, so a literal "# /// script" at column 0 inside the
# generator templates below would make uv see several metadata blocks in this
# file and refuse to run it.
PEP723_OPEN = "# " + "/// script"
PEP723_CLOSE = "# " + "///"


# =============================================================================
# Configuration - Pydantic Models for YAML Config
# =============================================================================

class ModelDefinition(BaseModel):
    """Definition of a single LLM model with provider configuration."""
    name: str
    provider: Literal["ollama", "anthropic", "openai", "google", "google-vertexai", "xai"]
    model: str
    endpoint: str | None = None
    api_key: str | None = None
    # Vertex AI specific fields
    credentials_file: str | None = None
    project: str | None = None
    location: str | None = None


class ConfigurationProfile(BaseModel):
    """Named configuration that assigns models to roles."""
    name: str
    description: str = ""
    ocr_model: str
    analysis_model: str
    organizer_model: str


class RenamerConfig(BaseModel):
    """Top-level configuration loaded from YAML."""
    models: list[ModelDefinition]
    configurations: list[ConfigurationProfile]
    default_configuration: str = "local-only"

    def get_model(self, name: str) -> ModelDefinition:
        """Get a model definition by name."""
        for model in self.models:
            if model.name == name:
                return model
        raise ValueError(f"Model '{name}' not found in configuration")

    def get_profile(self, name: str) -> ConfigurationProfile:
        """Get a configuration profile by name."""
        for profile in self.configurations:
            if profile.name == name:
                return profile
        raise ValueError(f"Configuration profile '{name}' not found")


def expand_env_vars(value: str) -> str:
    """Expand environment variables in the format ${VAR_NAME}."""
    pattern = re.compile(r'\$\{([^}]+)\}')

    def replacer(match: re.Match) -> str:
        var_name = match.group(1)
        env_value = os.environ.get(var_name)
        if env_value is None:
            raise ValueError(f"Environment variable '{var_name}' is not set")
        return env_value

    return pattern.sub(replacer, value)


def load_config(config_path: Path | None = None) -> RenamerConfig:
    """
    Load configuration from YAML file.

    Search order:
    1. Explicit path via --config flag
    2. ./renamer.yaml in current directory
    3. ~/.config/renamer/config.yaml

    Environment variables in ${VAR_NAME} format are expanded.
    """
    search_paths = []

    if config_path:
        search_paths.append(config_path)
    else:
        # Default search locations
        search_paths.append(Path.cwd() / "renamer.yaml")
        search_paths.append(Path.home() / ".config" / "renamer" / "config.yaml")

    config_file = None
    for path in search_paths:
        if path.exists():
            config_file = path
            break

    if config_file is None:
        searched = ", ".join(str(p) for p in search_paths)
        raise FileNotFoundError(
            f"No configuration file found. Searched: {searched}\n"
            "Create a config file or use --init-config to generate a template."
        )

    with open(config_file) as f:
        raw_config = yaml.safe_load(f)

    # Keep api_key values as-is during loading; they will be expanded lazily
    # when the model is actually used (in _create_model)

    return RenamerConfig(**raw_config)


def load_family_context(config_path: Path | None = None) -> str:
    """
    Load optional family.yaml file and format as context string.

    Search order:
    1. Same directory as config file (if provided)
    2. ./family.yaml in current directory
    3. ~/.config/renamer/family.yaml

    Returns empty string if file not found (this is not an error).
    """
    search_paths = []

    if config_path:
        # Look in same directory as config file
        search_paths.append(config_path.parent / "family.yaml")

    # Default search locations
    search_paths.append(Path.cwd() / "family.yaml")
    search_paths.append(Path.home() / ".config" / "renamer" / "family.yaml")

    family_file = None
    for path in search_paths:
        if path.exists():
            family_file = path
            break

    if family_file is None:
        return ""

    try:
        with open(family_file) as f:
            data = yaml.safe_load(f)

        if not data or "family" not in data:
            return ""

        # Format family members into readable context
        members = []
        for person in data["family"]:
            if person is None:
                continue
            # Handle both nested (name: {first: ...}) and flat ({name: null, first: ...}) structures
            name_info = person.get("name") if isinstance(person.get("name"), dict) else person
            parts = []

            first = name_info.get("first", "")
            middle = name_info.get("middle", "")
            last = name_info.get("last", "")
            nickname = name_info.get("nickname", "")

            full_name = " ".join(filter(None, [first, middle, last]))
            if full_name:
                parts.append(full_name)
            # Use nickname if provided, otherwise fall back to first name
            display_nickname = nickname or first
            if display_nickname:
                parts.append(f'(nickname: "{display_nickname}")')

            if parts:
                members.append(" ".join(parts))

        if not members:
            return ""

        return "Family members: " + "; ".join(members)

    except Exception as e:
        logger.warning(f"Could not load family.yaml: {e}")
        return ""


def generate_example_config() -> str:
    """Generate an example configuration file."""
    return '''# Renamer Configuration File
# Model definitions - each model has a unique name and provider config

models:
  # Local Ollama models
  - name: llava-local
    provider: ollama
    endpoint: http://localhost:11434
    model: llava:13b

  - name: llama32-local
    provider: ollama
    endpoint: http://localhost:11434
    model: llama3.2:latest

  - name: gemma3-local
    provider: ollama
    endpoint: http://localhost:11434
    model: gemma3:latest

  # Anthropic Claude
  - name: claude-sonnet
    provider: anthropic
    model: claude-sonnet-4-5-20250929
    api_key: ${ANTHROPIC_API_KEY}

  # OpenAI
  - name: gpt4o
    provider: openai
    model: gpt-4o
    api_key: ${OPENAI_API_KEY}

  # Google Gemini via API key (Google AI Studio)
  - name: gemini-flash
    provider: google
    model: gemini-2.0-flash
    api_key: ${GOOGLE_API_KEY}

  - name: gemini-pro
    provider: google
    model: gemini-2.5-pro
    api_key: ${GOOGLE_API_KEY}

  # Google Gemini via Vertex AI (service account credentials)
  - name: gemini3-flash-vertexai
    provider: google-vertexai
    model: gemini-3-flash-preview
    credentials_file: ./path/to/service-account.json
    project: your-gcp-project-id
    location: us-central1

  # xAI Grok (note: vision not supported via LangChain)
  - name: grok4
    provider: xai
    model: grok-4
    api_key: ${XAI_API_KEY}

# Named configurations that assign models to roles
configurations:
  - name: local-only
    description: All local models via Ollama
    ocr_model: llava-local
    analysis_model: llama32-local
    organizer_model: gemma3-local

  - name: hybrid-anthropic
    description: Local OCR, Anthropic analysis
    ocr_model: llava-local
    analysis_model: claude-sonnet
    organizer_model: claude-sonnet

  - name: hybrid-gemini
    description: Gemini for everything (excellent vision)
    ocr_model: gemini-flash
    analysis_model: gemini-flash
    organizer_model: gemini-pro

  - name: cloud-anthropic
    description: All Anthropic Claude
    ocr_model: claude-sonnet
    analysis_model: claude-sonnet
    organizer_model: claude-sonnet

  - name: budget-grok
    description: Grok for text analysis, local OCR (Grok vision not supported)
    ocr_model: llava-local
    analysis_model: grok4
    organizer_model: grok4

# Default configuration to use
default_configuration: local-only
'''


@dataclass
class Config:
    """Runtime application configuration (non-LLM settings)."""
    prompts_dir: Path = field(default_factory=lambda: Path(__file__).parent / "prompts")
    temp_dir: Path = field(default_factory=lambda: Path(tempfile.gettempdir()) / "renamer")
    output_dir: Path = field(default_factory=lambda: Path.home() / "Downloads")
    max_content_chars: int = 50000  # Max characters to send to LLM
    max_pages: int = 3  # Max pages to extract from documents

    # Supported file extensions
    supported_extensions: set[str] = field(default_factory=lambda: {
        ".pdf", ".md", ".txt",
        ".jpg", ".jpeg", ".png",
        ".docx", ".xlsx",
        ".key", ".numbers", ".pages",
    })


# =============================================================================
# Data Models
# =============================================================================

@dataclass
class FileRecord:
    """Record for a discovered file."""
    path: Path
    size: int
    modified: datetime
    extension: str

    def to_dict(self) -> dict:
        return {
            "path": str(self.path),
            "size": self.size,
            "modified": self.modified.isoformat(),
            "extension": self.extension,
        }

    @classmethod
    def from_dict(cls, data: dict) -> "FileRecord":
        return cls(
            path=Path(data["path"]),
            size=data["size"],
            modified=datetime.fromisoformat(data["modified"]),
            extension=data["extension"],
        )


@dataclass
class AnalysisResult:
    """Result from analyzing a single file."""
    original_path: str
    original_name: str
    ocr_model: str = ""
    ocr_time: str = "00:00.000"
    ocr_tokens_in: int = 0
    ocr_tokens_out: int = 0
    proposed_name: str = ""
    tags: list[str] = field(default_factory=list)
    summary: str = ""
    analysis_model: str = ""
    analysis_time: str = "00:00.000"
    analysis_tokens_in: int = 0
    analysis_tokens_out: int = 0
    error: str | None = None

    def to_dict(self) -> dict:
        return {
            "original_path": self.original_path,
            "original_name": self.original_name,
            "ocr_model": self.ocr_model,
            "ocr_time": self.ocr_time,
            "ocr_tokens_in": self.ocr_tokens_in,
            "ocr_tokens_out": self.ocr_tokens_out,
            "proposed_name": self.proposed_name,
            "tags": self.tags,
            "summary": self.summary,
            "analysis_model": self.analysis_model,
            "analysis_time": self.analysis_time,
            "analysis_tokens_in": self.analysis_tokens_in,
            "analysis_tokens_out": self.analysis_tokens_out,
            "error": self.error,
        }

    @classmethod
    def from_dict(cls, data: dict) -> "AnalysisResult":
        return cls(**data)


# =============================================================================
# macOS Native Dialogs (PyObjC)
# =============================================================================

class MacOSDialogs:
    """Native macOS dialogs using PyObjC."""

    @staticmethod
    def select_directory(title: str = "Select Directory to Scan") -> Path | None:
        """Show NSOpenPanel for directory selection."""
        try:
            from Cocoa import NSApp, NSApplication, NSModalResponseOK, NSOpenPanel

            # Ensure the application is properly initialized and activated
            # This is necessary when running from a terminal/CLI
            NSApplication.sharedApplication()
            NSApp.setActivationPolicy_(0)  # NSApplicationActivationPolicyRegular
            NSApp.activateIgnoringOtherApps_(True)

            panel = NSOpenPanel.openPanel()
            panel.setTitle_(title)
            panel.setCanChooseDirectories_(True)
            panel.setCanChooseFiles_(False)
            panel.setAllowsMultipleSelection_(False)
            panel.setCanCreateDirectories_(False)

            # Make the panel appear above other windows
            panel.setLevel_(3)  # NSFloatingWindowLevel

            # Run the panel
            if panel.runModal() == NSModalResponseOK:
                url = panel.URL()
                if url:
                    return Path(url.path())
            return None
        except ImportError:
            console.print("[yellow]PyObjC not available, falling back to input prompt[/yellow]")
            path_str = typer.prompt("Enter directory path")
            path = Path(path_str).expanduser().resolve()
            if path.is_dir():
                return path
            console.print(f"[red]Not a valid directory: {path}[/red]")
            return None

    @staticmethod
    def show_alert(title: str, message: str, buttons: list[str] | None = None) -> int:
        """Show NSAlert dialog, return button index (0 = first button)."""
        if buttons is None:
            buttons = ["OK"]
        try:
            from Cocoa import NSAlert, NSAlertFirstButtonReturn, NSApp, NSApplication

            # Ensure the application is properly initialized and activated
            NSApplication.sharedApplication()
            NSApp.setActivationPolicy_(0)  # NSApplicationActivationPolicyRegular
            NSApp.activateIgnoringOtherApps_(True)

            alert = NSAlert.alloc().init()
            alert.setMessageText_(title)
            alert.setInformativeText_(message)

            for button in buttons:
                alert.addButtonWithTitle_(button)

            result = alert.runModal()

            # Explicitly close the alert window and process the event to ensure it disappears
            window = alert.window()
            if window:
                window.orderOut_(None)
                window.close()
            # Process pending events to ensure the window is removed from screen
            NSApp.nextEventMatchingMask_untilDate_inMode_dequeue_(0xFFFFFFFF, None, "kCFRunLoopDefaultMode", False)

            return result - NSAlertFirstButtonReturn
        except ImportError:
            console.print(f"\n[bold]{title}[/bold]")
            console.print(message)
            for i, button in enumerate(buttons):
                console.print(f"  [{i}] {button}")
            choice = typer.prompt("Enter choice", default="0")
            return int(choice)

    @staticmethod
    def confirm(title: str, message: str) -> bool:
        """Show confirmation dialog."""
        result = MacOSDialogs.show_alert(title, message, ["Continue", "Cancel"])
        return result == 0

    @staticmethod
    def get_options_dialog() -> dict | None:
        """Show dialog to get processing options."""
        try:
            from Cocoa import (
                NSAlert,
                NSAlertFirstButtonReturn,
                NSApp,
                NSApplication,
                NSButton,
                NSMakeRect,
                NSOnState,
                NSTextField,
                NSView,
            )

            # Ensure the application is properly initialized and activated
            NSApplication.sharedApplication()
            NSApp.setActivationPolicy_(0)  # NSApplicationActivationPolicyRegular
            NSApp.activateIgnoringOtherApps_(True)

            alert = NSAlert.alloc().init()
            alert.setMessageText_("Processing Options")
            alert.setInformativeText_("Configure file processing settings:")
            alert.addButtonWithTitle_("Continue")
            alert.addButtonWithTitle_("Cancel")

            # Create accessory view for inputs
            accessory = NSView.alloc().initWithFrame_(NSMakeRect(0, 0, 300, 120))

            # Recurse checkbox
            recurse_check = NSButton.alloc().initWithFrame_(NSMakeRect(0, 90, 300, 20))
            recurse_check.setButtonType_(3)  # NSSwitchButton
            recurse_check.setTitle_("Process subdirectories")
            recurse_check.setState_(NSOnState)
            accessory.addSubview_(recurse_check)

            # Include pattern
            include_field = NSTextField.alloc().initWithFrame_(NSMakeRect(0, 60, 300, 24))
            include_field.setStringValue_(".*")
            include_field.setPlaceholderString_("Include pattern (regex)")
            accessory.addSubview_(include_field)

            # Exclude pattern
            exclude_field = NSTextField.alloc().initWithFrame_(NSMakeRect(0, 30, 300, 24))
            exclude_field.setStringValue_("")
            exclude_field.setPlaceholderString_("Exclude pattern (regex)")
            accessory.addSubview_(exclude_field)

            # Max files
            max_field = NSTextField.alloc().initWithFrame_(NSMakeRect(0, 0, 300, 24))
            max_field.setStringValue_("")
            max_field.setPlaceholderString_("Max files (leave empty for unlimited)")
            accessory.addSubview_(max_field)

            alert.setAccessoryView_(accessory)

            result = alert.runModal()

            # Explicitly close the alert window
            alert.window().orderOut_(None)

            if result == NSAlertFirstButtonReturn:
                max_str = max_field.stringValue()
                return {
                    "recurse": recurse_check.state() == NSOnState,
                    "include": include_field.stringValue() or ".*",
                    "exclude": exclude_field.stringValue() or "",
                    "max_files": int(max_str) if max_str else None,
                }
            return None
        except ImportError:
            # Fallback to terminal prompts
            console.print("\n[bold]Processing Options[/bold]")
            recurse = typer.confirm("Process subdirectories?", default=True)
            include = typer.prompt("Include pattern (regex)", default=".*")
            exclude = typer.prompt("Exclude pattern (regex, empty for none)", default="")
            max_str = typer.prompt("Max files (empty for unlimited)", default="")
            return {
                "recurse": recurse,
                "include": include,
                "exclude": exclude,
                "max_files": int(max_str) if max_str else None,
            }


# =============================================================================
# Temp File Manager
# =============================================================================

class TempFileManager:
    """Manage temporary files with Trash cleanup."""

    def __init__(self, config: Config):
        self.config = config
        self.temp_files: list[Path] = []
        self.session_dir: Path | None = None

    def initialize(self) -> Path:
        """Create session directory for temp files."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.session_dir = self.config.temp_dir / f"session_{timestamp}"
        self.session_dir.mkdir(parents=True, exist_ok=True)
        return self.session_dir

    def get_temp_path(self, filename: str) -> Path:
        """Get path for a temp file."""
        if not self.session_dir:
            self.initialize()
        assert self.session_dir is not None
        path = self.session_dir / filename
        self.temp_files.append(path)
        return path

    def cleanup(self):
        """Move all temp files to macOS Trash."""
        if self.session_dir and self.session_dir.exists():
            try:
                send2trash(str(self.session_dir))
                console.print(f"[dim]Moved temp files to Trash: {self.session_dir}[/dim]")
            except Exception as e:
                logger.warning(f"Could not move to trash: {e}")

    @staticmethod
    def append_jsonl(file_path: Path, record: dict):
        """Append JSON record to JSONL file."""
        with open(file_path, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")

    @staticmethod
    def read_jsonl(file_path: Path) -> Iterator[dict]:
        """Read JSONL file records."""
        if not file_path.exists():
            return
        with open(file_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line:
                    yield json.loads(line)


# =============================================================================
# Content Extractors
# =============================================================================

class ContentExtractor:
    """Extract text content from various file types."""

    def __init__(self, config: Config):
        self.config = config

    def extract(self, file_path: Path) -> tuple[str, list[bytes], bool]:
        """
        Extract text and images from file.
        Returns (text_content, list_of_image_bytes, needs_vision)
        """
        ext = file_path.suffix.lower()

        try:
            if ext == ".pdf":
                return self._extract_pdf(file_path)
            elif ext == ".docx":
                return self._extract_docx(file_path)
            elif ext == ".xlsx":
                return self._extract_xlsx(file_path)
            elif ext in {".txt", ".md"}:
                return self._extract_text(file_path)
            elif ext in {".jpg", ".jpeg", ".png"}:
                return self._extract_image(file_path)
            elif ext in {".key", ".numbers", ".pages"}:
                return self._extract_apple_format(file_path)
            else:
                return "", [], False
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            raise

    def _extract_pdf(self, path: Path) -> tuple[str, list[bytes], bool]:
        """Extract text from PDF, return images if scanned."""
        import fitz  # PyMuPDF

        doc = fitz.open(path)
        text_parts = []
        images = []
        has_text = False

        for page_num, page in enumerate(doc):
            if page_num >= self.config.max_pages:
                break

            # Try to extract text
            page_text = page.get_text()
            if page_text.strip():
                has_text = True
                text_parts.append(page_text)

            # If no text, render page as image for OCR
            if not page_text.strip():
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                images.append(pix.tobytes("png"))

        doc.close()

        text = "\n\n".join(text_parts)
        needs_vision = not has_text and len(images) > 0

        return text[:self.config.max_content_chars], images, needs_vision

    def _extract_docx(self, path: Path) -> tuple[str, list[bytes], bool]:
        """Extract text from Word document."""
        from docx import Document

        doc = Document(path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = "\n".join(paragraphs)
        return text[:self.config.max_content_chars], [], False

    def _extract_xlsx(self, path: Path) -> tuple[str, list[bytes], bool]:
        """Extract text from Excel spreadsheet."""
        from openpyxl import load_workbook

        wb = load_workbook(path, read_only=True, data_only=True)
        parts = []

        for sheet_name in wb.sheetnames[:3]:  # First 3 sheets
            sheet = wb[sheet_name]
            parts.append(f"=== Sheet: {sheet_name} ===")

            row_count = 0
            for row in sheet.iter_rows(max_row=50, values_only=True):  # First 50 rows
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    parts.append(row_text)
                    row_count += 1

        wb.close()
        text = "\n".join(parts)
        return text[:self.config.max_content_chars], [], False

    def _extract_text(self, path: Path) -> tuple[str, list[bytes], bool]:
        """Extract text from plain text files."""
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read(self.config.max_content_chars)
        return text, [], False

    def _extract_image(self, path: Path) -> tuple[str, list[bytes], bool]:
        """Return image bytes for vision processing."""
        with open(path, "rb") as f:
            image_bytes = f.read()
        return "", [image_bytes], True

    def _extract_apple_format(self, path: Path) -> tuple[str, list[bytes], bool]:
        """Export Apple format to PDF, then extract."""
        exporter = AppleScriptExporter()

        # Create temp PDF
        temp_pdf = Path(tempfile.gettempdir()) / f"temp_export_{path.stem}.pdf"

        try:
            if exporter.export_to_pdf(path, temp_pdf):
                text, images, needs_vision = self._extract_pdf(temp_pdf)
                return text, images, needs_vision
            else:
                logger.warning(f"Could not export {path}, returning empty content")
                return "", [], False
        finally:
            if temp_pdf.exists():
                temp_pdf.unlink()


# =============================================================================
# AppleScript Exporter
# =============================================================================

class AppleScriptExporter:
    """Export Apple iWork formats to PDF using AppleScript."""

    APP_MAP = {
        ".key": "Keynote",
        ".numbers": "Numbers",
        ".pages": "Pages",
    }

    def export_to_pdf(self, source_path: Path, output_path: Path) -> bool:
        """Export .key, .pages, .numbers to PDF."""
        ext = source_path.suffix.lower()
        app_name = self.APP_MAP.get(ext)

        if not app_name:
            return False

        script = f'''
        tell application "{app_name}"
            activate
            open POSIX file "{source_path}"
            delay 2
            set theDoc to front document
            export theDoc to POSIX file "{output_path}" as PDF
            close theDoc saving no
        end tell
        '''

        success, error = self._run_applescript(script)
        if not success:
            logger.warning(f"AppleScript export failed for {source_path}: {error}")
        return success

    def _run_applescript(self, script: str) -> tuple[bool, str]:
        """Execute AppleScript via osascript."""
        try:
            result = subprocess.run(
                ["osascript", "-e", script],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode != 0:
                return False, result.stderr
            return True, ""
        except subprocess.TimeoutExpired:
            return False, "AppleScript timed out"
        except Exception as e:
            return False, str(e)


# =============================================================================
# LLM Manager - Multi-Provider Support via LangChain
# =============================================================================

class LLMManager:
    """Manage LLM models across multiple providers using LangChain."""

    def __init__(self, renamer_config: RenamerConfig, profile_name: str | None = None):
        self.renamer_config = renamer_config
        self.profile_name = profile_name or renamer_config.default_configuration
        self.profile = renamer_config.get_profile(self.profile_name)
        self._models: dict[str, BaseChatModel] = {}
        self._ollama_client = None  # Lazy-loaded for Ollama-specific operations

    def _get_ollama_client(self):
        """Get or create Ollama client for model management."""
        if self._ollama_client is None:
            import ollama
            # Find first Ollama model to get endpoint
            for model_def in self.renamer_config.models:
                if model_def.provider == "ollama":
                    endpoint = model_def.endpoint or "http://localhost:11434"
                    self._ollama_client = ollama.Client(host=endpoint)
                    break
            if self._ollama_client is None:
                import ollama
                self._ollama_client = ollama.Client()
        return self._ollama_client

    def _resolve_api_key(self, model_def: ModelDefinition) -> str | None:
        """Resolve API key, expanding environment variables if needed."""
        if not model_def.api_key:
            return None

        api_key = model_def.api_key

        # Try to expand environment variable syntax ${VAR_NAME}
        if "${" in api_key:
            try:
                api_key = expand_env_vars(api_key)
            except ValueError:
                # If env var not set, try standard fallback env vars
                fallback_vars = {
                    "anthropic": "ANTHROPIC_API_KEY",
                    "openai": "OPENAI_API_KEY",
                    "google": "GOOGLE_API_KEY",
                    "xai": "XAI_API_KEY",
                }
                fallback = os.environ.get(fallback_vars.get(model_def.provider, ""))
                if fallback:
                    api_key = fallback
                else:
                    raise ValueError(
                        f"API key not set for model '{model_def.name}'. "
                        f"Set the environment variable or update your config."
                    )

        return api_key

    def _create_model(self, model_name: str) -> BaseChatModel:
        """Factory method to create LangChain model from config."""
        model_def = self.renamer_config.get_model(model_name)

        # Resolve API key lazily
        api_key = self._resolve_api_key(model_def)

        match model_def.provider:
            case "ollama":
                from langchain_ollama import ChatOllama
                return ChatOllama(
                    model=model_def.model,
                    base_url=model_def.endpoint or "http://localhost:11434",
                )
            case "anthropic":
                from langchain_anthropic import ChatAnthropic
                return ChatAnthropic(
                    model=model_def.model,
                    api_key=api_key,
                )
            case "openai":
                from langchain_openai import ChatOpenAI
                kwargs = {
                    "model": model_def.model,
                    "api_key": api_key,
                }
                if model_def.endpoint:
                    kwargs["base_url"] = model_def.endpoint
                return ChatOpenAI(**kwargs)
            case "google":
                from langchain_google_genai import ChatGoogleGenerativeAI
                return ChatGoogleGenerativeAI(
                    model=model_def.model,
                    google_api_key=api_key,
                )
            case "google-vertexai":
                from langchain_google_vertexai import ChatVertexAI
                kwargs = {"model": model_def.model}
                if model_def.project:
                    kwargs["project"] = model_def.project
                if model_def.location:
                    kwargs["location"] = model_def.location
                if model_def.credentials_file:
                    # Set the environment variable for Google credentials
                    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = model_def.credentials_file
                return ChatVertexAI(**kwargs)
            case "xai":
                from langchain_xai import ChatXAI
                return ChatXAI(
                    model=model_def.model,
                    xai_api_key=api_key,
                )
            case _:
                raise ValueError(f"Unknown provider: {model_def.provider}")

    def _get_or_create(self, model_name: str) -> BaseChatModel:
        """Get cached model or create new one."""
        if model_name not in self._models:
            self._models[model_name] = self._create_model(model_name)
        return self._models[model_name]

    @property
    def ocr_model(self) -> BaseChatModel:
        """Get the OCR/vision model."""
        return self._get_or_create(self.profile.ocr_model)

    @property
    def ocr_model_name(self) -> str:
        """Get the OCR model name for logging."""
        return self.profile.ocr_model

    @property
    def analysis_model(self) -> BaseChatModel:
        """Get the analysis model."""
        return self._get_or_create(self.profile.analysis_model)

    @property
    def analysis_model_name(self) -> str:
        """Get the analysis model name for logging."""
        return self.profile.analysis_model

    @property
    def organizer_model(self) -> BaseChatModel:
        """Get the organization model."""
        return self._get_or_create(self.profile.organizer_model)

    @property
    def organizer_model_name(self) -> str:
        """Get the organizer model name for logging."""
        return self.profile.organizer_model

    def _extract_token_usage(self, response) -> tuple[int, int]:
        """Extract token usage from LangChain response."""
        tokens_in = 0
        tokens_out = 0

        if hasattr(response, 'usage_metadata') and response.usage_metadata:
            tokens_in = response.usage_metadata.get('input_tokens', 0)
            tokens_out = response.usage_metadata.get('output_tokens', 0)
        elif hasattr(response, 'response_metadata'):
            metadata = response.response_metadata
            if 'usage' in metadata:
                usage = metadata['usage']
                tokens_in = usage.get('prompt_tokens', usage.get('input_tokens', 0))
                tokens_out = usage.get('completion_tokens', usage.get('output_tokens', 0))
            # Ollama format
            if 'prompt_eval_count' in metadata:
                tokens_in = metadata.get('prompt_eval_count', 0)
                tokens_out = metadata.get('eval_count', 0)

        return tokens_in, tokens_out

    def ensure_ollama_running(self) -> bool:
        """Check if Ollama is running (only relevant for Ollama models)."""
        # Check if any configured model uses Ollama
        ollama_models = [m for m in self.renamer_config.models if m.provider == "ollama"]
        if not ollama_models:
            return True  # No Ollama models, skip check

        try:
            client = self._get_ollama_client()
            client.list()
            return True
        except Exception:
            console.print("[yellow]Ollama not running, attempting to start...[/yellow]")
            try:
                subprocess.Popen(
                    ["ollama", "serve"],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
                time.sleep(3)
                client = self._get_ollama_client()
                client.list()
                console.print("[green]Ollama started successfully[/green]")
                return True
            except Exception as e:
                console.print(f"[red]Failed to start Ollama: {e}[/red]")
                console.print("[red]Please start Ollama manually with: ollama serve[/red]")
                return False

    def ensure_ollama_model(self, model_name: str) -> bool:
        """Pull Ollama model if not present."""
        model_def = self.renamer_config.get_model(model_name)
        if model_def.provider != "ollama":
            return True  # Not an Ollama model, skip

        try:
            client = self._get_ollama_client()
            models = client.list()
            model_names = [m.model for m in models.models] if models.models else []

            # Check if model exists (handle both with and without :latest suffix)
            base_name = model_def.model.split(":")[0]
            if not any(base_name in m for m in model_names):
                console.print(f"[yellow]Pulling model {model_def.model}...[/yellow]")
                with Progress(
                    SpinnerColumn(),
                    TextColumn("[progress.description]{task.description}"),
                    BarColumn(),
                    TaskProgressColumn(),
                    console=console,
                ) as progress:
                    task = progress.add_task(f"Downloading {model_def.model}", total=100)

                    for response in client.pull(model_def.model, stream=True):
                        if "completed" in response and "total" in response:
                            pct = (response["completed"] / response["total"]) * 100
                            progress.update(task, completed=pct)

                console.print(f"[green]Model {model_def.model} ready[/green]")
            return True
        except Exception as e:
            console.print(f"[red]Failed to pull model {model_def.model}: {e}[/red]")
            return False

    def ensure_models_ready(self) -> bool:
        """Ensure all models in the current profile are ready."""
        # Check Ollama is running if we have Ollama models
        if not self.ensure_ollama_running():
            return False

        # Ensure Ollama models are pulled
        for model_name in [self.profile.ocr_model, self.profile.analysis_model, self.profile.organizer_model]:
            if not self.ensure_ollama_model(model_name):
                return False

        return True

    def analyze_with_vision(self, images: list[bytes], prompt: str) -> tuple[str, int, int, float]:
        """
        Use vision model for OCR/image analysis.
        Returns (response_text, tokens_in, tokens_out, time_taken)
        """
        # Build multimodal message content
        content = [{"type": "text", "text": prompt}]

        for img in images:
            image_b64 = base64.b64encode(img).decode("utf-8")
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{image_b64}"}
            })

        message = HumanMessage(content=content)

        start_time = time.time()
        try:
            response = self.ocr_model.invoke([message])
            elapsed = time.time() - start_time

            # Handle response content - can be string or list of content blocks
            response_content = response.content
            if isinstance(response_content, list):
                # Extract text from content blocks (Gemini-style response)
                text = "".join(
                    block.get("text", "") if isinstance(block, dict) else str(block)
                    for block in response_content
                )
            else:
                text = response_content

            tokens_in, tokens_out = self._extract_token_usage(response)

            return text, tokens_in, tokens_out, elapsed
        except Exception as e:
            logger.error(f"Vision analysis failed: {e}")
            raise

    def analyze_document(self, text: str, filename: str, prompt_template: str) -> tuple[dict, int, int, float]:
        """
        Analyze document text, return structured data.
        Returns (parsed_response, tokens_in, tokens_out, time_taken)
        """
        # Fill in prompt template
        prompt = prompt_template.replace("{filename}", filename).replace("{content}", text)

        # Add JSON instruction to prompt
        json_prompt = prompt + "\n\nRespond with valid JSON only."

        message = HumanMessage(content=json_prompt)

        start_time = time.time()
        try:
            response = self.analysis_model.invoke([message])
            elapsed = time.time() - start_time

            # Handle response content - can be string or list of content blocks
            response_content = response.content
            if isinstance(response_content, list):
                # Extract text from content blocks (Gemini-style response)
                response_text = "".join(
                    block.get("text", "") if isinstance(block, dict) else str(block)
                    for block in response_content
                )
            else:
                response_text = response_content

            tokens_in, tokens_out = self._extract_token_usage(response)

            # Parse JSON response
            try:
                parsed = json.loads(response_text)
            except json.JSONDecodeError:
                # Try to extract JSON from response
                match = re.search(r'\{[^{}]*\}', response_text, re.DOTALL)
                if match:
                    parsed = json.loads(match.group())
                else:
                    parsed = {
                        "proposed_name": filename.rsplit(".", 1)[0],
                        "tags": ["unknown"],
                        "summary": "Could not analyze document.",
                    }

            return parsed, tokens_in, tokens_out, elapsed
        except Exception as e:
            logger.error(f"Document analysis failed: {e}")
            raise

    def generate_organization_plan(self, analysis_json: str, prompt_template: str) -> tuple[str, int, int, float]:
        """
        Generate folder structure markdown.
        Returns (response_text, tokens_in, tokens_out, time_taken)
        """
        prompt = prompt_template.replace("{analysis_json}", analysis_json)

        message = HumanMessage(content=prompt)

        start_time = time.time()
        try:
            response = self.organizer_model.invoke([message])
            elapsed = time.time() - start_time

            # Handle response content - can be string or list of content blocks
            response_content = response.content
            if isinstance(response_content, list):
                # Extract text from content blocks (Gemini-style response)
                text = "".join(
                    block.get("text", "") if isinstance(block, dict) else str(block)
                    for block in response_content
                )
            else:
                text = response_content

            tokens_in, tokens_out = self._extract_token_usage(response)

            return text, tokens_in, tokens_out, elapsed
        except Exception as e:
            logger.error(f"Organization planning failed: {e}")
            raise


# =============================================================================
# Pass 1: Discovery
# =============================================================================

class DiscoveryPass:
    """Pass 1: Discover and filter files."""

    def __init__(
        self,
        directory: Path,
        recurse: bool,
        include_pattern: str,
        exclude_pattern: str,
        max_files: int | None,
        config: Config,
        temp_manager: TempFileManager,
    ):
        self.directory = directory
        self.recurse = recurse
        self.include_pattern = re.compile(include_pattern) if include_pattern else None
        self.exclude_pattern = re.compile(exclude_pattern) if exclude_pattern else None
        self.max_files = max_files
        self.config = config
        self.temp_manager = temp_manager

    def run(self) -> Path:
        """Enumerate files, write to temp JSONL file. Returns path to temp file."""
        output_path = self.temp_manager.get_temp_path("pass1_discovery.jsonl")

        console.print(f"\n[bold]Pass 1: Discovering files in {self.directory}[/bold]")

        count = 0
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("Scanning...", total=None)

            for file_record in self._enumerate_files():
                if self.max_files and count >= self.max_files:
                    break

                TempFileManager.append_jsonl(output_path, file_record.to_dict())
                count += 1
                progress.update(task, description=f"Found {count} files...")

        console.print(f"[green]Discovered {count} files matching criteria[/green]")
        return output_path

    def _enumerate_files(self) -> Iterator[FileRecord]:
        """Walk directory tree, yield matching files."""
        if self.recurse:
            pattern = "**/*"
        else:
            pattern = "*"

        for path in self.directory.glob(pattern):
            if not path.is_file():
                continue

            # Skip hidden files
            if path.name.startswith("."):
                continue

            # Skip symlinks
            if path.is_symlink():
                continue

            # Check extension
            if path.suffix.lower() not in self.config.supported_extensions:
                continue

            # Apply include pattern
            if self.include_pattern and not self.include_pattern.search(path.name):
                continue

            # Apply exclude pattern
            if self.exclude_pattern and self.exclude_pattern.search(path.name):
                continue

            try:
                stat = path.stat()
                yield FileRecord(
                    path=path,
                    size=stat.st_size,
                    modified=datetime.fromtimestamp(stat.st_mtime),
                    extension=path.suffix.lower(),
                )
            except OSError as e:
                logger.warning(f"Could not stat {path}: {e}")


# =============================================================================
# Pass 2: Analysis
# =============================================================================

class AnalysisPass:
    """Pass 2: Analyze each file."""

    def __init__(
        self,
        discovery_file: Path,
        config: Config,
        temp_manager: TempFileManager,
        llm_manager: LLMManager,
        dry_run: bool = False,
        resume: bool = False,
        family_context: str = "",
    ):
        self.discovery_file = discovery_file
        self.config = config
        self.temp_manager = temp_manager
        self.llm = llm_manager
        self.dry_run = dry_run
        self.resume = resume
        self.family_context = family_context
        self.extractor = ContentExtractor(config)

        # Load prompts
        self.ocr_prompt = self._load_prompt("ocr_vision.txt")
        self.analysis_prompt = self._load_prompt("document_analysis.txt")

        # Inject family context into analysis prompt if available
        if self.family_context:
            self.analysis_prompt = self._inject_family_context(self.analysis_prompt)

    def _inject_family_context(self, prompt: str) -> str:
        """Inject family context into prompt before the document content."""
        context_block = f"\nCONTEXT:\n{self.family_context}\n"
        # Insert before "DOCUMENT FILENAME:" if present, otherwise append
        if "DOCUMENT FILENAME:" in prompt:
            return prompt.replace("DOCUMENT FILENAME:", context_block + "DOCUMENT FILENAME:")
        return prompt + context_block

    def _load_prompt(self, filename: str) -> str:
        """Load prompt from file."""
        prompt_path = self.config.prompts_dir / filename
        if prompt_path.exists():
            return prompt_path.read_text()
        else:
            logger.warning(f"Prompt file not found: {prompt_path}")
            return ""

    def run(self) -> Path:
        """Analyze each file, return path to analysis JSONL."""
        output_path = self.temp_manager.get_temp_path("pass2_analysis.jsonl")
        error_log_path = self.temp_manager.get_temp_path("errors.log")

        # Load already processed files for resume
        processed_paths = set()
        if self.resume and output_path.exists():
            for record in TempFileManager.read_jsonl(output_path):
                processed_paths.add(record["original_path"])
            console.print(f"[dim]Resuming: {len(processed_paths)} files already processed[/dim]")

        # Count total files
        total = sum(1 for _ in TempFileManager.read_jsonl(self.discovery_file))

        console.print(f"\n[bold]Pass 2: Analyzing {total} files[/bold]")

        if self.dry_run:
            console.print("[yellow]Dry run mode - skipping LLM analysis[/yellow]")

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TaskProgressColumn(),
            console=console,
        ) as progress:
            task = progress.add_task("Analyzing...", total=total)

            for i, record_dict in enumerate(TempFileManager.read_jsonl(self.discovery_file)):
                record = FileRecord.from_dict(record_dict)

                # Skip if already processed (resume mode)
                if str(record.path) in processed_paths:
                    progress.update(task, advance=1)
                    continue

                progress.update(
                    task,
                    description=f"[{i+1}/{total}] {record.path.name[:40]}..."
                )

                result = self._analyze_file(record)
                TempFileManager.append_jsonl(output_path, result.to_dict())

                if result.error:
                    with open(error_log_path, "a") as f:
                        f.write(f"{record.path}: {result.error}\n")

                progress.update(task, advance=1)

        console.print(f"[green]Analysis complete. Results saved to {output_path}[/green]")
        return output_path

    def _analyze_file(self, record: FileRecord) -> AnalysisResult:
        """Process a single file."""
        result = AnalysisResult(
            original_path=str(record.path),
            original_name=record.path.name,
        )

        if self.dry_run:
            result.proposed_name = record.path.stem
            result.tags = ["dry-run"]
            result.summary = "Dry run - no analysis performed"
            return result

        try:
            # Extract content
            text, images, needs_vision = self.extractor.extract(record.path)

            # OCR if needed
            if needs_vision and images:
                try:
                    ocr_text, ocr_in, ocr_out, ocr_time = self.llm.analyze_with_vision(
                        images[:self.config.max_pages],
                        self.ocr_prompt,
                    )
                    text = ocr_text
                    result.ocr_model = self.llm.ocr_model_name
                    result.ocr_time = f"{int(ocr_time // 60):02d}:{ocr_time % 60:06.3f}"
                    result.ocr_tokens_in = ocr_in
                    result.ocr_tokens_out = ocr_out

                    # Write OCR text back to PDF if it's a PDF
                    if record.extension == ".pdf":
                        self._add_ocr_layer_to_pdf(record.path, ocr_text)
                except Exception as e:
                    logger.error(f"OCR failed for {record.path}: {e}")
                    result.error = f"OCR failed: {e}"
                    return result

            # Analyze document
            if not text.strip():
                result.error = "No text content extracted"
                return result

            try:
                parsed, analysis_in, analysis_out, analysis_time = self.llm.analyze_document(
                    text,
                    record.path.name,
                    self.analysis_prompt,
                )

                result.proposed_name = parsed.get("proposed_name", record.path.stem)
                result.tags = parsed.get("tags", [])
                result.summary = parsed.get("summary", "")
                result.analysis_model = self.llm.analysis_model_name
                result.analysis_time = f"{int(analysis_time // 60):02d}:{analysis_time % 60:06.3f}"
                result.analysis_tokens_in = analysis_in
                result.analysis_tokens_out = analysis_out

            except Exception as e:
                logger.error(f"Analysis failed for {record.path}: {e}")
                result.error = f"Analysis failed: {e}"

            return result

        except Exception as e:
            logger.error(f"Error processing {record.path}: {e}")
            result.error = str(e)
            return result

    def _add_ocr_layer_to_pdf(self, pdf_path: Path, ocr_text: str):
        """Add OCR text layer to PDF."""
        import fitz

        # Check if file is writable
        if not os.access(pdf_path, os.W_OK):
            # Create a copy with -ocr suffix
            new_path = pdf_path.with_stem(pdf_path.stem + "-ocr")
            console.print(f"[dim]Creating OCR copy: {new_path.name}[/dim]")

            import shutil
            shutil.copy2(pdf_path, new_path)
            pdf_path = new_path

        try:
            doc = fitz.open(pdf_path)

            # Add invisible text layer to first page
            if len(doc) > 0:
                page = doc[0]
                # Insert text as invisible annotation
                rect = page.rect
                page.insert_text(
                    (rect.x0 + 10, rect.y0 + 10),
                    ocr_text[:1000],  # Limit text length
                    fontsize=1,  # Very small
                    color=(1, 1, 1),  # White (invisible on white background)
                )

            doc.saveIncr()
            doc.close()
        except Exception as e:
            logger.warning(f"Could not add OCR layer to {pdf_path}: {e}")


# =============================================================================
# Pass 3: Organization
# =============================================================================

class OrganizationPass:
    """Pass 3: Generate folder structure and scripts."""

    def __init__(
        self,
        analysis_file: Path,
        base_directory: Path,
        config: Config,
        temp_manager: TempFileManager,
        llm_manager: LLMManager,
        dry_run: bool = False,
        family_context: str = "",
    ):
        self.analysis_file = analysis_file
        self.base_directory = base_directory
        self.config = config
        self.temp_manager = temp_manager
        self.llm = llm_manager
        self.dry_run = dry_run
        self.family_context = family_context

        self.org_prompt = self._load_prompt("organization_planning.txt")

        # Inject family context into organization prompt if available
        if self.family_context:
            self.org_prompt = self._inject_family_context(self.org_prompt)

    def _inject_family_context(self, prompt: str) -> str:
        """Inject family context into prompt before the analysis data."""
        context_block = f"\nCONTEXT:\n{self.family_context}\n"
        # Insert before "ANALYZED DOCUMENTS:" if present, otherwise append
        if "ANALYZED DOCUMENTS:" in prompt:
            return prompt.replace("ANALYZED DOCUMENTS:", context_block + "ANALYZED DOCUMENTS:")
        return prompt + context_block

    def _load_prompt(self, filename: str) -> str:
        """Load prompt from file."""
        prompt_path = self.config.prompts_dir / filename
        if prompt_path.exists():
            return prompt_path.read_text()
        return ""

    def run(self) -> tuple[Path | None, Path | None, Path | None]:
        """Generate organization plan and scripts."""
        console.print("\n[bold]Pass 3: Generating organization plan[/bold]")

        # Load all analysis records
        records = list(TempFileManager.read_jsonl(self.analysis_file))

        if not records:
            console.print("[red]No analysis records found[/red]")
            return None, None, None

        # Filter out records with errors
        valid_records = [r for r in records if not r.get("error")]
        console.print(f"[dim]Processing {len(valid_records)} successfully analyzed files[/dim]")

        if self.dry_run:
            console.print("[yellow]Dry run mode - generating placeholder plan[/yellow]")
            plan_text = "# Dry Run\n\nNo actual organization plan generated in dry run mode."
            file_mappings = []
        else:
            # Generate organization plan with LLM
            analysis_json = json.dumps(valid_records, indent=2)

            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                console=console,
            ) as progress:
                task = progress.add_task("Generating organization plan...", total=None)

                plan_text, _, _, elapsed = self.llm.generate_organization_plan(
                    analysis_json,
                    self.org_prompt,
                )

                progress.update(task, description=f"Plan generated in {elapsed:.1f}s")

            # Parse file mappings from plan
            file_mappings = self._extract_file_mappings(plan_text)

        # Generate output files
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # Organization plan markdown
        plan_path = self.config.output_dir / f"file_organization_plan_{timestamp}.md"
        plan_path.write_text(plan_text)
        console.print(f"[green]Plan saved: {plan_path}[/green]")

        # Apply script
        apply_path = self.config.output_dir / f"apply_changes_{timestamp}.py"
        apply_script = self._generate_apply_script(file_mappings)
        apply_path.write_text(apply_script)
        console.print(f"[green]Apply script saved: {apply_path}[/green]")

        # Revert script
        revert_path = self.config.output_dir / f"revert_changes_{timestamp}.py"
        revert_script = self._generate_revert_script(file_mappings)
        revert_path.write_text(revert_script)
        console.print(f"[green]Revert script saved: {revert_path}[/green]")

        return plan_path, apply_path, revert_path

    def _extract_file_mappings(self, plan_text: str) -> list[dict]:
        """Extract JSON file mappings from plan text."""
        try:
            # Find JSON array in the response
            match = re.search(r'\[\s*\{[^]]+\}\s*\]', plan_text, re.DOTALL)
            if match:
                return json.loads(match.group())
        except json.JSONDecodeError:
            pass

        logger.warning("Could not parse file mappings from plan")
        return []

    def _generate_apply_script(self, mappings: list[dict]) -> str:
        """Generate Python script to apply changes."""
        moves_str = json.dumps(mappings, indent=4)
        base_dir = str(self.base_directory)

        return f'''#!/usr/bin/env python3
{PEP723_OPEN}
# requires-python = ">=3.10"
# dependencies = []
{PEP723_CLOSE}
"""
Generated by renamer.py on {datetime.now().isoformat()}
Applies file organization changes.

Run with: uv run apply_changes_*.py
"""
import json
import shutil
from pathlib import Path

BASE_DIR = Path("{base_dir}")

MAPPINGS = {moves_str}

def main():
    log_file = Path(__file__).with_suffix(".log")
    operations = []

    for mapping in MAPPINGS:
        original = Path(mapping["original_path"])
        dest_folder = BASE_DIR / mapping["destination_folder"]
        new_name = mapping["new_filename"]
        new_path = dest_folder / new_name

        if not original.exists():
            print(f"SKIP: {{original}} (not found)")
            continue

        # Handle filename conflicts
        if new_path.exists():
            stem = new_path.stem
            suffix = new_path.suffix
            counter = 1
            while new_path.exists():
                new_path = dest_folder / f"{{stem}}-{{counter}}{{suffix}}"
                counter += 1

        # Create destination directory
        dest_folder.mkdir(parents=True, exist_ok=True)

        # Move file
        shutil.move(str(original), str(new_path))
        print(f"MOVED: {{original}} -> {{new_path}}")

        operations.append({{
            "original": str(original),
            "new": str(new_path),
        }})

    # Save log for revert
    with open(log_file, "w") as f:
        json.dump(operations, f, indent=2)

    print(f"\\nOperations logged to: {{log_file}}")
    print(f"Total files moved: {{len(operations)}}")

if __name__ == "__main__":
    main()
'''

    def _generate_revert_script(self, mappings: list[dict]) -> str:
        """Generate Python script to revert changes."""
        return f'''#!/usr/bin/env python3
{PEP723_OPEN}
# requires-python = ">=3.10"
# dependencies = ["send2trash>=1.8.0"]
{PEP723_CLOSE}
"""
Generated by renamer.py on {datetime.now().isoformat()}
Reverts file organization changes.

Run with: uv run revert_changes_*.py
"""
import json
import shutil
from pathlib import Path
from send2trash import send2trash

def main():
    # Find the apply log file
    script_path = Path(__file__)
    log_file = script_path.parent / script_path.name.replace("revert_", "apply_").replace(".py", ".log")

    if not log_file.exists():
        print(f"ERROR: Log file not found: {{log_file}}")
        print("Make sure you ran the apply script first.")
        return

    with open(log_file) as f:
        operations = json.load(f)

    # Reverse operations
    created_dirs = set()
    for op in reversed(operations):
        current = Path(op["new"])
        original = Path(op["original"])

        if not current.exists():
            print(f"SKIP: {{current}} (not found)")
            continue

        # Restore original directory if needed
        original.parent.mkdir(parents=True, exist_ok=True)

        # Move back
        shutil.move(str(current), str(original))
        print(f"REVERTED: {{current}} -> {{original}}")

        # Track parent directories for cleanup
        created_dirs.add(current.parent)

    # Clean up empty directories
    for dir_path in sorted(created_dirs, reverse=True):
        try:
            if dir_path.exists() and not any(dir_path.iterdir()):
                send2trash(str(dir_path))
                print(f"REMOVED: {{dir_path}}")
        except Exception as e:
            print(f"Could not remove {{dir_path}}: {{e}}")

    # Move log to trash
    send2trash(str(log_file))
    print(f"\\nRevert complete. Log moved to trash.")

if __name__ == "__main__":
    main()
'''


# =============================================================================
# Main Application
# =============================================================================

@app.command()
def init_config(
    output: Annotated[
        Path | None,
        typer.Argument(help="Output path for config file"),
    ] = None,
):
    """Generate an example configuration file."""
    if output is None:
        output = Path.cwd() / "renamer.yaml"

    if output.exists():
        console.print(f"[red]File already exists: {output}[/red]")
        raise typer.Exit(1)

    output.write_text(generate_example_config())
    console.print(f"[green]Configuration file created: {output}[/green]")
    console.print("[dim]Edit this file to configure your LLM providers and profiles.[/dim]")


@app.command()
def run(
    directory: Annotated[
        Path | None,
        typer.Argument(help="Directory to scan (opens file picker if not provided)"),
    ] = None,
    recurse: Annotated[
        bool,
        typer.Option("--recurse/--no-recurse", help="Process subdirectories"),
    ] = True,
    include: Annotated[
        str,
        typer.Option("--include", "-i", help="Include regex pattern"),
    ] = ".*",
    exclude: Annotated[
        str,
        typer.Option("--exclude", "-e", help="Exclude regex pattern"),
    ] = "",
    max_files: Annotated[
        int | None,
        typer.Option("--max-files", "-n", help="Maximum files to process"),
    ] = None,
    dry_run: Annotated[
        bool,
        typer.Option("--dry-run", help="Skip LLM calls, show what would be processed"),
    ] = False,
    auto_continue: Annotated[
        bool,
        typer.Option("--auto-continue", help="Skip confirmation prompts"),
    ] = False,
    resume: Annotated[
        bool,
        typer.Option("--resume", help="Resume from previous interrupted run"),
    ] = False,
    config_path: Annotated[
        Path | None,
        typer.Option("--config", "-c", help="Path to YAML configuration file"),
    ] = None,
    profile: Annotated[
        str | None,
        typer.Option("--profile", "-p", help="Configuration profile to use"),
    ] = None,
):
    """Run the file renamer and organizer."""

    # Load configuration from YAML
    try:
        renamer_config = load_config(config_path)
    except FileNotFoundError as e:
        console.print(f"[red]{e}[/red]")
        console.print("[yellow]Run 'renamer.py init-config' to create a configuration file.[/yellow]")
        raise typer.Exit(1)
    except Exception as e:
        console.print(f"[red]Error loading configuration: {e}[/red]")
        raise typer.Exit(1)

    # Initialize runtime config (non-LLM settings)
    config = Config()

    # Load optional family context
    family_context = load_family_context(config_path)
    if family_context:
        console.print(f"[dim]Loaded family context: {family_context}[/dim]")

    # Initialize temp manager
    temp_manager = TempFileManager(config)
    temp_manager.initialize()

    # Initialize LLM manager
    try:
        llm_manager = LLMManager(renamer_config, profile)
    except ValueError as e:
        console.print(f"[red]Configuration error: {e}[/red]")
        raise typer.Exit(1)

    console.print(Panel.fit(
        "[bold blue]File Renamer & Organizer[/bold blue]\n"
        f"Profile: {llm_manager.profile_name}",
        border_style="blue",
    ))

    try:
        # Get directory
        if directory is None:
            directory = MacOSDialogs.select_directory()
            if directory is None:
                console.print("[red]No directory selected. Exiting.[/red]")
                return
        else:
            directory = directory.expanduser().resolve()

        if not directory.is_dir():
            console.print(f"[red]Not a valid directory: {directory}[/red]")
            raise typer.Exit(1)

        console.print(f"[dim]Target directory: {directory}[/dim]")

        # Get options via dialog if not all provided via CLI
        if not all([include != ".*", exclude, max_files]):
            if not auto_continue:
                options = MacOSDialogs.get_options_dialog()
                if options is None:
                    console.print("[red]Cancelled by user.[/red]")
                    return
                recurse = options["recurse"]
                include = options["include"]
                exclude = options["exclude"]
                max_files = options["max_files"]

        if not dry_run:
            # Ensure models are available
            console.print("\n[bold]Checking models...[/bold]")
            if not llm_manager.ensure_models_ready():
                raise typer.Exit(1)

        # Pass 1: Discovery
        discovery = DiscoveryPass(
            directory=directory,
            recurse=recurse,
            include_pattern=include,
            exclude_pattern=exclude,
            max_files=max_files,
            config=config,
            temp_manager=temp_manager,
        )
        discovery_file = discovery.run()

        # Show discovered files
        records = list(TempFileManager.read_jsonl(discovery_file))
        if not records:
            console.print("[yellow]No files found matching criteria.[/yellow]")
            return

        # Display file list
        table = Table(title="Discovered Files")
        table.add_column("File", style="cyan")
        table.add_column("Size", justify="right")
        table.add_column("Modified")

        for record in records[:20]:  # Show first 20
            path = Path(record["path"])
            size = record["size"]
            size_str = f"{size:,}" if size < 1024 else f"{size/1024:.1f}K" if size < 1024*1024 else f"{size/1024/1024:.1f}M"
            table.add_row(
                str(path.relative_to(directory)),
                size_str,
                record["modified"][:10],
            )

        if len(records) > 20:
            table.add_row(f"... and {len(records) - 20} more files", "", "")

        console.print(table)
        console.print(f"\n[dim]Full list saved to: {discovery_file}[/dim]")

        # Confirm to continue
        if not auto_continue:
            if not MacOSDialogs.confirm(
                "Continue with Analysis?",
                f"Found {len(records)} files. Proceed to analyze with LLM?"
            ):
                console.print("[yellow]Cancelled by user.[/yellow]")
                return

        # Pass 2: Analysis
        analysis = AnalysisPass(
            discovery_file=discovery_file,
            config=config,
            temp_manager=temp_manager,
            llm_manager=llm_manager,
            dry_run=dry_run,
            resume=resume,
            family_context=family_context,
        )
        analysis_file = analysis.run()

        # Pass 3: Organization
        organization = OrganizationPass(
            analysis_file=analysis_file,
            base_directory=directory,
            config=config,
            temp_manager=temp_manager,
            llm_manager=llm_manager,
            dry_run=dry_run,
            family_context=family_context,
        )
        plan_path, apply_path, revert_path = organization.run()

        # Final summary
        console.print("\n" + "=" * 60)
        console.print("[bold green]Processing Complete![/bold green]")
        console.print("=" * 60)
        console.print("\nOutput files saved to Downloads:")
        if plan_path:
            console.print(f"  - Organization plan: {plan_path.name}")
        if apply_path:
            console.print(f"  - Apply script: {apply_path.name}")
        if revert_path:
            console.print(f"  - Revert script: {revert_path.name}")
        console.print(f"\nTo apply changes: [cyan]uv run {apply_path}[/cyan]")
        console.print(f"To revert changes: [cyan]uv run {revert_path}[/cyan]")

    except KeyboardInterrupt:
        console.print("\n[yellow]Interrupted by user.[/yellow]")
    except Exception as e:
        console.print(f"\n[red]Error: {e}[/red]")
        logger.exception("Unexpected error")
        raise typer.Exit(1)
    finally:
        # Cleanup temp files
        temp_manager.cleanup()


if __name__ == "__main__":
    app()
